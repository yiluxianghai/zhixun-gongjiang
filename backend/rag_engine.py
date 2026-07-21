"""
RAG引擎 - 文档处理与检索增强生成
工程监理巡视问题分类闭环管理智能体

功能：
1. PDF/DOCX文档文本提取
2. 文本分块（支持重叠）
3. 向量嵌入生成（OpenAI API / 关键词降级）
4. 语义检索（余弦相似度 / 关键词匹配降级）
5. RAG上下文构建
"""
import os
import io
import json
import re
from typing import Optional, List

import httpx

# ========== 文本提取 ==========

def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件提取文本"""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """从DOCX文件提取文本"""
    import docx
    doc = docx.Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                text_parts.append(row_text)
    return "\n\n".join(text_parts)


def extract_text(file_path: str, file_type: str) -> str:
    """根据文件类型提取文本"""
    file_type = file_type.lower()
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_path)
    elif file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


# ========== 文本分块 ==========

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 200) -> List[str]:
    """将文本分块（带重叠）
    
    Args:
        text: 原始文本
        chunk_size: 每块最大字符数
        overlap: 相邻块重叠字符数
    
    Returns:
        分块列表
    """
    if not text or not text.strip():
        return []
    
    # 清理文本
    text = re.sub(r'\n{3,}', '\n\n', text)  # 多个空行合并
    text = text.strip()
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        if end >= len(text):
            chunks.append(text[start:].strip())
            break
        
        # 尝试在句号/换行处切分，避免截断句子
        cut_point = end
        for i in range(end, max(end - 100, start), -1):
            if text[i] in '。.！!？?；;\n':
                cut_point = i + 1
                break
        
        chunk = text[start:cut_point].strip()
        if chunk:
            chunks.append(chunk)
        
        start = cut_point - overlap if cut_point - overlap > start else cut_point
    
    return chunks


# ========== 向量嵌入 ==========

async def generate_embedding(text: str, api_key: str, base_url: str) -> Optional[List[float]]:
    """调用OpenAI兼容API生成文本嵌入向量
    
    Args:
        text: 要嵌入的文本
        api_key: API密钥
        base_url: API基础URL
    
    Returns:
        嵌入向量（float列表），失败返回None
    """
    if not api_key or not text.strip():
        return None
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{base_url}/embeddings",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "text-embedding-3-small",
                    "input": text[:8000],  # 截断超长文本
                }
            )
            response.raise_for_status()
            result = response.json()
            return result["data"][0]["embedding"]
    except Exception as e:
        print(f"[RAG] 嵌入向量生成失败: {e}")
        return None


# ========== 关键词检索（降级方案） ==========

def keyword_search(query: str, chunks: List[str], top_k: int = 5) -> List[str]:
    """基于关键词匹配的文本检索（无嵌入向量时的降级方案）
    
    使用改进的TF-IDF相似度计算
    """
    if not chunks:
        return []
    
    # 分词（简单的中文分词：按字+双字组合）
    def get_terms(text: str) -> set:
        text = text.lower()
        terms = set()
        # 单字
        for c in text:
            if c.strip() and c not in '，。、；：！？\n\r\t 。，;:!?()（）[]【】""\'\'':
                terms.add(c)
        # 双字组合
        for i in range(len(text) - 1):
            bi = text[i:i+2]
            if bi.strip() and not all(c in '，。、；：！？\n\r\t 。，;:!?()（）[]【】""\'\'' for c in bi):
                terms.add(bi)
        return terms
    
    query_terms = get_terms(query)
    if not query_terms:
        return chunks[:top_k]
    
    scores = []
    for idx, chunk in enumerate(chunks):
        chunk_terms = get_terms(chunk)
        if not chunk_terms:
            scores.append((idx, 0))
            continue
        # Jaccard相似度
        intersection = len(query_terms & chunk_terms)
        union = len(query_terms | chunk_terms)
        score = intersection / union if union > 0 else 0
        scores.append((idx, score))
    
    # 按分数排序，取top_k
    scores.sort(key=lambda x: x[1], reverse=True)
    top_indices = [idx for idx, score in scores[:top_k] if score > 0]
    
    if not top_indices:
        # 如果没有匹配，返回前top_k个
        return chunks[:top_k]
    
    return [chunks[idx] for idx in top_indices]


# ========== 语义检索 ==========

async def retrieve_context(
    query: str,
    chunks_data: List[dict],  # [{"content": str, "embedding": [float]}]
    api_key: str = "",
    base_url: str = "",
    top_k: int = 5,
) -> List[str]:
    """RAG检索：从文档分块中检索与查询最相关的内容
    
    Args:
        query: 查询文本（问题描述）
        chunks_data: 分块数据列表
        api_key: API密钥（用于嵌入查询）
        base_url: API基础URL
        top_k: 返回最相关的top_k个分块
    
    Returns:
        相关文本块列表
    """
    if not chunks_data:
        return []
    
    # 尝试向量检索
    if api_key:
        query_embedding = await generate_embedding(query, api_key, base_url)
        if query_embedding:
            try:
                import numpy as np
                
                # 构建嵌入矩阵
                embeddings = []
                valid_chunks = []
                for cd in chunks_data:
                    if cd.get("embedding"):
                        emb = json.loads(cd["embedding"])
                        embeddings.append(emb)
                        valid_chunks.append(cd["content"])
                
                if embeddings:
                    query_vec = np.array(query_embedding)
                    emb_matrix = np.array(embeddings)
                    # 余弦相似度
                    similarities = np.dot(emb_matrix, query_vec) / (
                        np.linalg.norm(emb_matrix, axis=1) * np.linalg.norm(query_vec) + 1e-8
                    )
                    # 取top_k
                    top_indices = np.argsort(similarities)[::-1][:top_k]
                    result = [valid_chunks[i] for i in top_indices if similarities[i] > 0.1]
                    if result:
                        print(f"[RAG] 向量检索完成，返回 {len(result)} 个相关块")
                        return result
            except Exception as e:
                print(f"[RAG] 向量检索失败，降级为关键词检索: {e}")
    
    # 降级：关键词检索
    all_chunks = [cd["content"] for cd in chunks_data]
    result = keyword_search(query, all_chunks, top_k)
    print(f"[RAG] 关键词检索完成，返回 {len(result)} 个相关块")
    return result


# ========== 文档处理流水线 ==========

async def process_document(
    file_path: str,
    file_type: str,
    doc_id: int,
    kb_id: int,
    db_session,
    api_key: str = "",
    base_url: str = "",
):
    """文档处理流水线：提取 → 分块 → 嵌入 → 存储
    
    Args:
        file_path: 文件路径
        file_type: 文件类型
        doc_id: 文档数据库ID
        kb_id: 知识库ID
        db_session: 数据库会话
        api_key: API密钥（用于生成嵌入向量）
        base_url: API基础URL
    """
    from database import KnowledgeDocument, DocumentChunk
    
    try:
        # 1. 提取文本
        print(f"[RAG] 开始处理文档 (doc_id={doc_id}): {file_path}")
        text = extract_text(file_path, file_type)
        print(f"[RAG] 文本提取完成，长度: {len(text)} 字符")
        
        if not text.strip():
            doc = db_session.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
            if doc:
                doc.status = "error"
                doc.error_message = "文档内容为空或无法提取文本"
                db_session.commit()
            return
        
        # 2. 分块
        chunks = chunk_text(text, chunk_size=800, overlap=200)
        print(f"[RAG] 文本分块完成: {len(chunks)} 个块")
        
        # 3. 更新文档状态
        doc = db_session.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if doc:
            doc.text_content = text
            doc.chunk_count = len(chunks)
        
        # 4. 删除旧分块（如果存在）
        db_session.query(DocumentChunk).filter(DocumentChunk.doc_id == doc_id).delete()
        
        # 5. 为每个块生成嵌入向量并存储
        for idx, chunk in enumerate(chunks):
            embedding_str = ""
            
            # 尝试生成嵌入向量
            if api_key:
                emb = await generate_embedding(chunk, api_key, base_url)
                if emb:
                    embedding_str = json.dumps(emb)
            
            chunk_record = DocumentChunk(
                doc_id=doc_id,
                chunk_index=idx,
                content=chunk,
                embedding=embedding_str,
            )
            db_session.add(chunk_record)
        
        # 6. 更新文档状态为完成
        if doc:
            doc.status = "ready"
            doc.error_message = ""
        db_session.commit()
        print(f"[RAG] 文档处理完成 (doc_id={doc_id}): {len(chunks)} 个块已存储")
        
    except Exception as e:
        print(f"[RAG] 文档处理失败 (doc_id={doc_id}): {e}")
        doc = db_session.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if doc:
            doc.status = "error"
            doc.error_message = str(e)
            db_session.commit()


# ========== RAG上下文构建 ==========

def build_rag_context(retrieved_chunks: List[str], max_chars: int = 3000) -> str:
    """将检索到的文本块构建为LLM可用的上下文文本"""
    if not retrieved_chunks:
        return ""
    
    parts = []
    total = 0
    for i, chunk in enumerate(retrieved_chunks):
        if total + len(chunk) > max_chars:
            # 截断超长块
            remaining = max_chars - total
            if remaining > 100:
                parts.append(f"[片段{i+1}]\n{chunk[:remaining]}...")
            break
        parts.append(f"[片段{i+1}]\n{chunk}")
        total += len(chunk)
    
    return "\n\n".join(parts)
